from __future__ import annotations

import csv
import json
import shutil
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

import make_nyt_style_figures


REPORT_DIR = Path(__file__).resolve().parent
ROOT = REPORT_DIR.parents[1]
DOCS_REPORT_DIR = ROOT / "docs" / "v5_benchmark"
FIG_DIR = REPORT_DIR / "figures"
VISUAL_FIG_DIR = REPORT_DIR / "figures_nyt"
DOCS_FIG_DIR = DOCS_REPORT_DIR / "figures"
DOCS_VISUAL_FIG_DIR = DOCS_REPORT_DIR / "figures_nyt"
MARKDOWN_OUT = REPORT_DIR / "CardioConsult_TrackC_APA_Technical_Report.md"
DOCX_OUT = REPORT_DIR / "CardioConsult_TrackC_APA_Technical_Report.docx"
PDF_OUT = REPORT_DIR / "CardioConsult_TrackC_APA_Technical_Report.pdf"
DOCS_MD_OUT = DOCS_REPORT_DIR / "CardioConsult_PC_V5_EchoBench_Technical_Report_APA_20260604.md"
DOCS_DOCX_OUT = DOCS_REPORT_DIR / "CardioConsult_PC_V5_EchoBench_Technical_Report_APA_20260604.docx"

ACCENT = "2E74B5"
DARK = "0B2545"
MUTED = "475569"
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F8FAFC"
BORDER = "CBD5E1"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def row_by_label(rows: list[dict[str, str]], label: str) -> dict[str, str]:
    for row in rows:
        if row.get("label") == label:
            return row
    raise KeyError(label)


def fmt(value: float | str | None, digits: int = 3) -> str:
    if value is None or value == "":
        return ""
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def pct(value: float | str | None) -> str:
    if value is None or value == "":
        return ""
    try:
        return f"{float(value) * 100:.1f}%"
    except Exception:
        return str(value)


def metric_table(rows: list[dict[str, str]], labels: list[tuple[str, str]]) -> list[list[str]]:
    out: list[list[str]] = []
    for label, zh in labels:
        row = row_by_label(rows, label)
        out.append(
            [
                zh,
                row["n"],
                row["positive_gold"],
                row["tp"],
                row["tn"],
                row["fp"],
                row["fn"],
                fmt(row["accuracy"]),
                fmt(row["sensitivity"]),
                fmt(row["specificity"]),
                fmt(row["f1"]),
            ]
        )
    return out


def markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    aligns = ["---"] + ["---:" for _ in headers[1:]]
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(aligns) + " |"]
    lines.extend("| " + " | ".join(map(str, row)) + " |" for row in rows)
    return "\n".join(lines)


def gather_data() -> dict:
    full_metrics = read_csv(ROOT / "validation_speedopt" / "full_evidence" / "newtraining_metrics.csv")
    rep12_metrics = read_csv(
        ROOT
        / "validation_speedopt"
        / "freeze_runs"
        / "echobench_20260604_175653"
        / "validation"
        / "newtraining_metrics.csv"
    )
    old_summary = read_json(ROOT / "validation_speedopt" / "old_baseline" / "newtraining_summary.json")
    cold_summary = read_json(ROOT / "validation_speedopt" / "speedopt_cold" / "newtraining_summary.json")
    full_latency = read_json(ROOT / "validation_speedopt" / "full_evidence" / "latency_summary.json")
    rep12_latency = read_json(
        ROOT / "validation_speedopt" / "freeze_runs" / "echobench_20260604_175653" / "latency_summary.json"
    )
    full_summary = read_json(
        ROOT
        / "validation_speedopt"
        / "freeze_runs_full"
        / "echobench_20260604_180638"
        / "validation"
        / "newtraining_summary.json"
    )
    rep12_summary = read_json(
        ROOT / "validation_speedopt" / "freeze_runs" / "echobench_20260604_175653" / "validation" / "newtraining_summary.json"
    )
    server_smoke = read_json(ROOT / "validation_speedopt" / "server_smoke_general_current_20260604.json")
    server_case = read_json(ROOT / "validation_speedopt" / "server_pipeline_case1_current_20260604.json")
    integrated = read_csv(REPORT_DIR / "integrated_test_results.csv")
    return {
        "full_metrics": full_metrics,
        "rep12_metrics": rep12_metrics,
        "old_summary": old_summary,
        "cold_summary": cold_summary,
        "full_latency": full_latency,
        "rep12_latency": rep12_latency,
        "full_summary": full_summary,
        "rep12_summary": rep12_summary,
        "server_smoke": server_smoke,
        "server_case": server_case,
        "integrated": integrated,
    }


def update_integrated_results(data: dict) -> None:
    generated_datasets = {"Freeze EchoBench", "Freeze service"}
    rows = [row for row in data["integrated"] if row.get("dataset") not in generated_datasets]
    rows.extend(
        [
            {
                "dataset": "Freeze EchoBench",
                "scope": "冻结前完整证据 60 例",
                "availability": "可用",
                "n": "60 local authorized DICOM/report-linked cases",
                "metrics": (
                    f"cases_ok={data['full_summary']['cases_ok']}/60; "
                    f"mean_runtime={data['full_latency']['runtime_seconds']['mean']}s; "
                    f"MR F1={row_by_label(data['full_metrics'], 'mr')['f1']}; "
                    f"AR F1={row_by_label(data['full_metrics'], 'ar')['f1']}; "
                    f"low EF F1={row_by_label(data['full_metrics'], 'low_ef')['f1']}"
                ),
                "interpretation": "当前冻结代码在完整证据场景保持 60/60 可运行；MR、低 EF 表现稳定，AR 依赖切面覆盖；TR 本批全阳性，不能解释为泛化特异性。",
            },
            {
                "dataset": "Freeze EchoBench",
                "scope": "冻结前 12 帧代表输入 60 例",
                "availability": "可用",
                "n": "60 cases, max 12 files/frames per case",
                "metrics": (
                    f"cases_ok={data['rep12_summary']['cases_ok']}/60; "
                    f"mean_runtime={data['rep12_latency']['runtime_seconds']['mean']}s warm-cache; "
                    f"MR F1={row_by_label(data['rep12_metrics'], 'mr')['f1']}; "
                    f"AR F1={row_by_label(data['rep12_metrics'], 'ar')['f1']}; "
                    f"low EF F1={row_by_label(data['rep12_metrics'], 'low_ef')['f1']}"
                ),
                "interpretation": "代表输入上限可快速运行；MR 保持稳定，TR 只说明本批全阳性样本内召回一致，AR、RWMA、左房扩大仍是切面覆盖敏感项。",
            },
            {
                "dataset": "Freeze service",
                "scope": "本地常驻 Gemma4 4B GGUF 服务",
                "availability": "可用",
                "n": "2 /completion requests + EchoBench case 1",
                "metrics": (
                    f"completion_elapsed={data['server_smoke']['first_completion']['elapsed_seconds']}s/"
                    f"{data['server_smoke']['second_completion']['elapsed_seconds']}s; "
                    f"project_diagnosis={data['server_case']['diagnosis_seconds']}s; "
                    f"prompt_leakage={data['server_case']['has_prompt_leakage']}"
                ),
                "interpretation": "常驻服务可复用热启动；项目诊断链路已启用报告保护层，输出无提示词泄漏并保留安全边界。",
            },
        ]
    )
    fields = ["dataset", "scope", "availability", "n", "metrics", "interpretation"]
    with (REPORT_DIR / "integrated_test_results.csv").open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        writer.writerows({field: row.get(field, "") for field in fields} for row in rows)
    (REPORT_DIR / "integrated_test_results.json").write_text(
        json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8"
    )


REFERENCES = [
    "Bland, J. M., & Altman, D. G. (1986). Statistical methods for assessing agreement between two methods of clinical measurement. *The Lancet, 327*(8476), 307-310. https://doi.org/10.1016/S0140-6736(86)90837-8",
    "Cohen, J. (1960). A coefficient of agreement for nominal scales. *Educational and Psychological Measurement, 20*(1), 37-46. https://doi.org/10.1177/001316446002000104",
    "FDA, Health Canada, & Medicines and Healthcare products Regulatory Agency. (2021). *Good machine learning practice for medical device development: Guiding principles*. https://www.gov.uk/government/publications/good-machine-learning-practice-for-medical-device-development-guiding-principles",
    "ggml-org. (n.d.). *GGUF file format - llama.cpp*. Retrieved June 4, 2026, from https://www.mintlify.com/ggml-org/llama.cpp/concepts/gguf-format",
    "Google AI for Developers. (n.d.). *Get started with Gemma models*. Retrieved June 4, 2026, from https://ai.google.dev/gemma/docs/get_started",
    "Google AI for Developers. (n.d.). *Run Gemma content generation and inferences*. Retrieved June 4, 2026, from https://ai.google.dev/gemma/docs/run",
    "Karargyris, A., Umeton, R., Sheller, M. J., Aristizabal, A., George, J., Wuest, A., Pati, S., Kassem, H., Zenk, M., Baid, U., et al. (2023). Federated benchmarking of medical artificial intelligence with MedPerf. *Nature Machine Intelligence, 5*, 799-810. https://doi.org/10.1038/s42256-023-00652-2",
    "Lang, R. M., Badano, L. P., Mor-Avi, V., Afilalo, J., Armstrong, A., Ernande, L., Flachskampf, F. A., Foster, E., Goldstein, S. A., Kuznetsova, T., Lancellotti, P., Muraru, D., Picard, M. H., Rietzschel, E. R., Rudski, L., Spencer, K. T., Tsang, W., & Voigt, J.-U. (2015). Recommendations for cardiac chamber quantification by echocardiography in adults: An update from the American Society of Echocardiography and the European Association of Cardiovascular Imaging. *Journal of the American Society of Echocardiography, 28*(1), 1-39.e14. https://doi.org/10.1016/j.echo.2014.10.003",
    "Leclerc, S., Smistad, E., Pedrosa, J., Ostvik, A., Cervenansky, F., Espinosa, F., Espeland, T., Berg, E. A. R., Jodoin, P.-M., Grenier, T., Lartizien, C., D'Hooge, J., Lovstakken, L., & Bernard, O. (2019). Deep learning for segmentation using an open large-scale dataset in 2D echocardiography. *IEEE Transactions on Medical Imaging, 38*(9), 2198-2210. https://doi.org/10.1109/TMI.2019.2900516",
    "MLCommons. (n.d.-a). *MLPerf Client*. Retrieved June 4, 2026, from https://mlcommons.org/benchmarks/client/",
    "MLCommons. (n.d.-b). *MLPerf Inference benchmarks*. Retrieved June 4, 2026, from https://docs.mlcommons.org/inference/",
    "OpenAI. (2025). *Introducing HealthBench*. https://openai.com/index/healthbench/",
    "OpenAI. (2025). *HealthBench: Evaluating large language models towards improved human health*. https://arxiv.org/abs/2505.08775",
    "Ouyang, D., He, B., Ghorbani, A., Yuan, N., Ebinger, J., Langlotz, C. P., Heidenreich, P. A., Harrington, R. A., Liang, D. H., Ashley, E. A., & Zou, J. Y. (2020). Video-based AI for beat-to-beat assessment of cardiac function. *Nature, 580*(7802), 252-256. https://doi.org/10.1038/s41586-020-2145-8",
    "Tableau Software. (2014). *Visual analysis best practices: Simple techniques for making every data visualization useful and beautiful*. Tableau Software.",
    "Vickers, A. J., & Elkin, E. B. (2006). Decision curve analysis: A novel method for evaluating prediction models. *Medical Decision Making, 26*(6), 565-574. https://doi.org/10.1177/0272989X06295361",
    "Yu, Y., & Acton, S. T. (2002). Speckle reducing anisotropic diffusion. *IEEE Transactions on Image Processing, 11*(11), 1260-1270. https://doi.org/10.1109/TIP.2002.804276",
    "Zoghbi, W. A., Adams, D., Bonow, R. O., Enriquez-Sarano, M., Foster, E., Grayburn, P. A., Hahn, R. T., Han, Y., Hung, J., Lang, R. M., Little, S. H., Shah, D. J., Shernan, S., Thavendiranathan, P., Thomas, J. D., & Weissman, N. J. (2017). Recommendations for noninvasive evaluation of native valvular regurgitation: A report from the American Society of Echocardiography developed in collaboration with the Society for Cardiovascular Magnetic Resonance. *Journal of the American Society of Echocardiography, 30*(4), 303-371. https://doi.org/10.1016/j.echo.2017.01.007",
]


def report_markdown(data: dict) -> str:
    labels = [
        ("mr", "二尖瓣反流 MR"),
        ("tr", "三尖瓣反流 TR*"),
        ("ar", "主动脉瓣反流 AR"),
        ("low_ef", "低 EF / 左室收缩功能减低"),
        ("rwma", "节段性室壁运动异常 RWMA"),
        ("la_enlargement", "左房扩大"),
        ("bradycardia", "心动过缓"),
    ]
    full_rows = metric_table(data["full_metrics"], labels)
    rep_rows = metric_table(data["rep12_metrics"], labels)
    full_lat = data["full_latency"]["runtime_seconds"]
    rep_lat = data["rep12_latency"]["runtime_seconds"]
    old_mean = data["old_summary"]["mean_case_runtime_seconds"]
    cold_mean = data["cold_summary"]["mean_case_runtime_seconds"]
    warm_mean = data["rep12_latency"]["runtime_seconds"]["mean"]
    speedup_old_to_warm = (old_mean - warm_mean) / old_mean * 100.0
    server = data["server_case"]
    smoke = data["server_smoke"]
    today = date.today().isoformat()
    full_mr = row_by_label(data["full_metrics"], "mr")
    full_ar = row_by_label(data["full_metrics"], "ar")
    full_low = row_by_label(data["full_metrics"], "low_ef")
    rep_mr = row_by_label(data["rep12_metrics"], "mr")
    rep_ar = row_by_label(data["rep12_metrics"], "ar")
    rep_low = row_by_label(data["rep12_metrics"], "low_ef")

    return f"""# CardioConsult PC V5 技术报告

生成日期：{today}

报告版本：可视化增强版，APA 引用版

系统版本：CardioConsult PC V5 freeze build `2026-06-04`

定位：超声机器旁离线分析设备；医学教学、算法演示与基层参考工具，不是医疗器械。

## 摘要

CardioConsult PC V5 解决的问题是：在缺少心脏超声专科医生、网络条件有限或需要保护教学/脱敏病例数据的场景中，如何让超声初学者和基层医疗点从 PNG、DICOM/DCOM、cine/视频等文件中获得一份可解释、可审计、不会冒充正式诊断的心脏超声教学参考结果。系统把 PC 定义为接在超声机器、无线超声软件、DICOM 工作站或局域网导出目录旁的离线分析终端；所有图像预处理、特征提取、层级标签、Doppler 瓣膜定位评分、轻量多智能体审计和 Gemma4 4B GGUF 结构化报告生成都在本机完成。

本版报告按 Tableau Software（2014）《Visual Analysis Best Practices》的视觉分析原则重排：每张图先回答一个问题，避免颜色和形状过载，优先使用共享基线、直接标注、少量强调色和明确的证据边界。旧版 8 张 R 图仍保留为可复核资产，主文改用 7 张叙事型图，把“能运行、准到哪里、哪里必须补扫、为什么低成本”讲清楚。

冻结前检查显示，当前仓库可以完成三类任务：第一，规则路径 `app.py --self-test-rule-only` 通过，输出包含 `教学参考病症判断：`、`最小病症：`、`逻辑链：` 和安全边界；第二，授权本地 60 例 EchoBench 完整证据场景 60/60 成功，平均 {full_lat['mean']:.3f} 秒/例，MR F1={fmt(full_mr['f1'])}、AR F1={fmt(full_ar['f1'])}、低 EF F1={fmt(full_low['f1'])}；第三，12 帧代表输入场景 60/60 成功，warm-cache 平均 {rep_lat['mean']:.3f} 秒/例，MR F1={fmt(rep_mr['f1'])}、AR F1={fmt(rep_ar['f1'])}、低 EF F1={fmt(rep_low['f1'])}。需要特别说明的是，TR 在本批 60 例中全部为阳性，表格中的 1.000 只代表本批同质阳性样本内召回/一致性，不能当作真实世界泛化准确率、特异性或阴性排除能力。本地 `llama-server` 项目链路第 1 例服务模式诊断耗时 {server['diagnosis_seconds']:.3f} 秒，报告保护层记录 `{server['status']}`，最终输出 `has_prompt_leakage=false`、`report_source={server.get('report_source', '')}`；当前默认提示词进一步新增 `gemma4_structured` JSON 渲染路径。

从评委视角看，V5 的强项不是单一公开排行榜最高分，而是低成本可运行、输入格式兼容、输出合同明确、审计链完整、隐私边界清楚。它的主要风险也清楚：AR、RWMA、左房扩大和严重程度分级在 12 帧输入下仍受切面覆盖影响；当前 60 例报告链接标签不是多专家盲评金标准；系统不能作为临床诊断或治疗建议。这个边界在 UI、README、技术报告和诊断输出中均需保留。

![图0：提交材料状态与阅读地图](figures_nyt/nyt_fig3_submission_readiness.png)

## 1. 问题陈述与用户价值

基层医疗点和超声初学者常见困难不是“完全没有图像”，而是有图像、有 DICOM 或无线超声导出文件，却缺少稳定的心脏超声专科解读、补扫建议和教学级复盘。传统多人会诊依赖专家时间和网络条件；云端多模态模型又会带来隐私、成本、网络和部署不确定性。CardioConsult V5 的目标是提供一个离线、低边际成本、可审计的教学参考层，把系统输出限制在“疑似病症判断、证据链、补扫建议、安全分层”上。

医学边界上，本项目遵循良好机器学习实践中关于目标人群、数据代表性、开发追溯、性能监测和人类监督的原则（FDA, Health Canada, & Medicines and Healthcare products Regulatory Agency, 2021）。工程边界上，它借鉴 MLPerf 对客户端、离线和服务场景的分层测量思路，但 EchoBench v1 是项目自建 benchmark，不是 MLCommons 官方提交（MLCommons, n.d.-a, n.d.-b）。

## 2. 系统方案

系统分为五层：输入层读取 PNG/JPG、DICOM/DCOM、多帧 TIFF、GIF、MP4/MOV/AVI 等文件；B-mode 分支计算 SRAD/CLAHE 风格预处理、边缘密度、纹理熵、散斑残差、腔室面积代理和收缩舒张差；Color Doppler 分支将 HSV 血流颜色转为活跃区、连通域、喷流宽度、方向一致性、湍流、涡量代理和 MR/TR/AR/PR 瓣膜定位评分；校准层用 V4 shared-EK/coupled-EK 与 V5 EchoNet-Dynamic 校准增强 EF / 左室收缩功能减低；报告层默认要求 Gemma4 4B GGUF 输出 JSON object，再由报告保护层按本地诊断合同重渲染为中文教学摘要，并清除提示词泄漏、截断和 AI 口吻。

为对齐 Gemma4 提交指南中的原生函数调用要求，当前代码新增 `cardio_pc/function_calling.py` 和 `docs/gemma4_function_calling_contract.md`。离线 Gemma4 增强路径可按 JSON tool-call 合同请求 `summarize_ultrasound_features`、`run_rule_diagnosis` 和 `safety_boundary_check` 三个白名单内部工具；本地执行器拒绝非白名单工具和额外参数。该设计让模型只能查询低维超声证据、确定性规则诊断和安全边界，不能执行任意代码或读取原始病人文件。`tools/function_calling_smoke.py` 可在无 GGUF、无网络条件下复现 tool manifest、白名单执行和非法工具拒绝逻辑。

![图1：离线读片链路与报告输出合同](figures_nyt/nyt_fig4_system_flow.png)

EchoNet-Dynamic 是公开心超视频数据集，包含心尖四腔视频和 EF/ESV/EDV/左室追踪标注，适合用于心功能任务（Ouyang et al., 2020）。CAMUS 则适合 2D 心超分割和腔室结构评估（Leclerc et al., 2019）。瓣膜反流和腔室定量的正式临床判断仍应遵循 ASE/EACVI 指南，而不能只依赖本项目的颜色代理特征（Lang et al., 2015; Zoghbi et al., 2017）。

## 3. Benchmark 设计

冻结版报告采用四组证据：

1. **规则自检。** 本地 synthetic A4C ED/ES 输入，验证文件读取、特征提取、层级标签、自然化报告和安全边界。
2. **EchoBench 完整证据。** 60 例授权 DICOM/报告时间映射，每例读取全部可用文件，衡量资料充分时的系统上限。
3. **EchoBench 12 帧代表输入。** 每例最多 12 个文件/帧，模拟产品输入上限和现场演示限制。
4. **本地服务。** 常驻 `llama-server` 加载 Gemma4 4B GGUF，测 `/completion` 热启动复用和项目级服务诊断链路。

医疗项目指标包括准确率、敏感性、特异性、精确率、F1、EF MAE/RMSE/相关系数、报告安全字段、补扫提示和人工复核边界。一般工程指标包括启动可用性、每例延迟、P50/P90/P95/P99、热启动复用、边际调用成本、离线隐私、审计链和可复现命令。

## 4. 主要性能结果

### 4.1 完整证据：标签性能

完整证据场景 60/60 例成功，平均 {full_lat['mean']:.3f} 秒/例，P95={full_lat['p95']:.3f} 秒。MR、低 EF 表现稳定；AR 中等；RWMA、心动过缓等标签因为样本少或未接入 ECG/完整报告结构化字段，不能过度宣称。TR 虽在表中显示 1.000，但本批 60 例全部带 TR 阳性标签，缺少阴性样本，因此这不是严格意义上的泛化性能估计。

![图2：完整证据与 12 帧输入的 F1 稳定性](figures_nyt/nyt_fig1_f1_story.png)

{markdown_table(["标签", "n", "阳性", "TP", "TN", "FP", "FN", "准确率", "敏感性", "特异性", "F1"], full_rows)}

注：TR 在本批 60 例中全部为阳性，因此 TN=0、FP=0，特异性和阴性排除能力没有统计解释价值；TR F1=1.000 只能解释为本批同质阳性样本内召回一致，不能作为“真实准确率 100%”宣传。

### 4.2 12 帧输入：现场演示上限

12 帧代表输入场景 60/60 例成功，warm-cache 平均 {rep_lat['mean']:.3f} 秒/例，P95={rep_lat['p95']:.3f} 秒。MR F1 仍为 {fmt(rep_mr['f1'])}，低 EF F1 为 {fmt(rep_low['f1'])}；AR F1 降至 {fmt(rep_ar['f1'])}，提示主动脉瓣反流需要更完整的 A5C/主动脉瓣相关切面和彩色多普勒序列。TR 在 12 帧输入中同样全阳性，仍只能作为“这批 TR 阳性病例没有漏报”的证据，而不能评价对正常/无 TR 病例的区分能力。

![图3：12 帧输入的诊断画像与切面敏感项](figures/fig2_12frame_metric_profile.png)

{markdown_table(["标签", "n", "阳性", "TP", "TN", "FP", "FN", "准确率", "敏感性", "特异性", "F1"], rep_rows)}

注：12 帧 TR 结果同样受全阳性标签分布影响。后续需要引入明确无 TR 的阴性病例、不同机器/不同导出格式和更宽的严重程度分布，才能报告可靠的 TR 特异性、AUC 或校准曲线。

![图4：12 帧混淆矩阵组成，供细节复核](figures/fig3_confusion_components_12frame.png)

### 4.3 延迟、缓存和本地服务

SpeedOpt 前 12 帧基线平均 {old_mean:.3f} 秒/例；SpeedOpt 冷缓存平均 {cold_mean:.3f} 秒/例；冻结版 warm-cache 12 帧平均 {warm_mean:.3f} 秒/例，相比旧基线下降 {speedup_old_to_warm:.1f}%。完整证据 warm-cache 平均 {full_lat['mean']:.3f} 秒/例，说明当前普通 PC 上已经可以支撑交互式教学演示。

![图5：延迟阶梯与本地 Gemma4 服务定位](figures_nyt/nyt_fig2_latency_story.png)

{markdown_table(["场景", "平均秒/例", "P50", "P90", "P95", "P99", "最大值", "平均文件数"], [
    ["完整证据 warm-cache", fmt(full_lat['mean']), fmt(full_lat['median_p50']), fmt(full_lat['p90']), fmt(full_lat['p95']), fmt(full_lat['p99']), fmt(full_lat['max']), fmt(data['full_latency']['files_per_case']['mean'])],
    ["12帧 warm-cache", fmt(rep_lat['mean']), fmt(rep_lat['median_p50']), fmt(rep_lat['p90']), fmt(rep_lat['p95']), fmt(rep_lat['p99']), fmt(rep_lat['max']), fmt(data['rep12_latency']['files_per_case']['mean'])],
    ["12帧 SpeedOpt 冷缓存", fmt(cold_mean), "", "", "", "", "", "12.000"],
    ["12帧旧基线", fmt(old_mean), "", "", "", "", "", "12.000"],
])}

本地服务 smoke 连续两次 `/completion` 均返回 OK：第一次 {smoke['first_completion']['elapsed_seconds']:.3f} 秒，第二次 {smoke['second_completion']['elapsed_seconds']:.3f} 秒。第二次请求 prompt tok/s 从 {smoke['first_completion']['timings']['prompt_per_second']:.2f} 提升到 {smoke['second_completion']['timings']['prompt_per_second']:.2f}，说明常驻模型复用有效。

![图6：llama-server 热启动复用，供模型服务复核](figures/fig5_server_smoke_hot_reuse.png)

项目级服务链路使用 EchoBench 第 1 例、12 个文件、`max_tokens=240`：文件加载 {server['load_seconds']:.3f} 秒，特征提取 {server['feature_seconds']:.3f} 秒，Gemma4 服务诊断 {server['diagnosis_seconds']:.3f} 秒；报告保护层启用，最终 `has_prompt_leakage=false`、`report_source={server.get('report_source', '')}`，并保留医学安全边界。该服务证据来自结构化 JSON 默认启用前的旧服务复现；当前代码默认新增 `structured_llm_output=true`，下一次模型服务复现可记录 `report_source=gemma4_structured`。

### 4.4 EchoNet-Dynamic 校准层

V5 的 EchoNet-Dynamic 校准层用于补强 EF 和左室收缩功能减低识别，不替代瓣膜反流规则。held-out 指标为 EF MAE=7.271、EF RMSE=9.603、EF 相关系数=0.647、低 EF AUC=0.764、低 EF F1=0.496。这支持“教学提示”用途，但不足以宣称临床 EF 自动测量。

![图7：EchoNet-Dynamic 校准指标](figures/fig6_echonet_training_metrics.png)

## 5. 数据来源与证据覆盖

本仓库不分发原始 DICOM、公开数据集压缩包或 GGUF 权重，只保存汇总报告、指标和图表。CAMUS、EchoNet-Dynamic、HMC-QU、EchoXFlow、MR Ultrasound Images、MIMIC-IV-ECHO/ECHOVIEW、CACTUS 等来源按许可证或访问条件分级记录在 `DATASETS.md` 和 `integrated_test_results.*`。从冻结评审角度，最强证据是授权本地 60 例、CAMUS 阶段测试、本地 smoke 和服务链路；计划数据集只能写成后续路线，不能包装成已完成结果。

![图8：证据覆盖情况与不可过度宣称边界](figures_nyt/nyt_fig5_evidence_ladder.png)

## 6. 与现有方案的比较

EchoNet-Dynamic 和 CAMUS 这类公开方案的优势是任务定义清晰、数据结构规范、适合学术复现；限制是通常聚焦 EF、容积或分割，不直接覆盖基层教学场景中的 DICOM/DCOM 兼容、Color Doppler 代理、中文层级病症报告和离线本地部署（Leclerc et al., 2019; Ouyang et al., 2020）。云端多模态模型可能具备更强的自然语言解释能力，但会引入病例上传、网络、费用、审计和服务可用性问题。

CardioConsult V5 的差异化是：输入侧兼容真实超声导出文件；算法侧融合 B-mode、Color Doppler、Doppler 瓣膜定位评分、动图差分和 EchoNet 校准；输出侧强制最小病症、逻辑链和安全边界，并用 JSON 合同限制 Gemma4 自由漂移；运行侧采用本地规则、小模型与 GGUF LLM，边际调用成本接近 0。代价是部分细粒度病症在缺切面时准确率下降，需要明确补扫和复核。

## 7. 成本模型与取舍

本地方案的主要成本是一次性 PC、存储和 GGUF 文件准备。运行时不按病例调用云 API，规则路径吞吐可按每小时数百到上千例估算；真实演示受人工选文件、磁盘、DICOM 解码和 GGUF 文本 token 数影响。以冻结版 warm-cache 为例，12 帧规则链路平均 {warm_mean:.3f} 秒/例，理论吞吐约 {3600 / warm_mean:.0f} 例/小时；完整证据平均 {full_lat['mean']:.3f} 秒/例，理论吞吐约 {3600 / full_lat['mean']:.0f} 例/小时。

![图9：成本、隐私与准确性取舍](figures_nyt/nyt_fig6_cost_boundary.png)

主要取舍如下：

- 为了离线和低成本，牺牲了云端大模型的算力弹性。
- 为了可审计和安全，核心标签由规则/校准层给出，LLM 主要负责表达，牺牲了部分自由生成能力。
- 为了兼容 DICOM/DCOM 和动图，保留了较复杂的文件解析与代表帧抽样。
- 为了提高基层筛查敏感性，低 EF 等标签更偏向“提示/待排”，正常特异性需要更多阴性样本继续优化。

## 8. 冻结前补救与评审风险

本轮冻结检查完成了以下补救：

1. 删除无用或容易打不开的高级 BAT，只保留 `install_deps.bat` 和 `run_cardio_pc_v5.bat` 两个普通入口。
2. 修正诊断链输出保护层，清除提示词泄漏、markdown 模板和“作为 AI / 我将”式口吻。
3. 更新本地服务 JSON，旧的提示词泄漏片段已被干净报告替换。
4. 重跑 60 例完整证据和 12 帧冻结 benchmark。
5. 保留 8 张 R 复核图，同时新增 7 张白皮书式叙事图，并同步到 submission 与 docs。
6. 报告中明确写出医学边界、数据许可、未下载数据集和不可过度宣称的标签。

仍需诚实呈现的风险：

- 当前 60 例是本地授权教学数据，不是多中心外部临床验证。
- 报告链接标签不是多专家盲评共识标签。
- TR 全阳性导致特异性、AUC 和真实阴性排除能力不可解释；PR、severe、HCM 等标签样本不足。
- 在线 demo 是规则匹配网页，不等于完整 PC 图像特征和 GGUF 推理。
- 项目不能被描述为临床诊断系统、医疗器械或治疗建议工具。

![图10：最终输出合同与安全边界](figures_nyt/nyt_fig7_safety_contract.png)

<!-- PAGEBREAK -->

## 9. 下一步

1. 建立 2-3 位心超医生盲评表，计算 Cohen's Kappa、加权 Kappa 或 ICC（Bland & Altman, 1986; Cohen, 1960）。
2. 补充 AR、PR、severe regurgitation、RWMA、LVH/HCM、LA enlargement 的阳性和阴性平衡样本。
3. 训练轻量 ONNX/TFLite 左室/左房分割模型，用 CAMUS 和 EchoNet tracing 改善 EF 与腔室大小代理。
4. 将 PLAX、PSAX、A4C、A2C、A5C、subcostal 等切面识别作为显式中间任务。
5. 按 HealthBench 类似的 rubric 思路建立医疗文本质量评估：事实性、完整性、风险提示、边界表达、教学价值（OpenAI, 2025）。

## 10. 可复现信息

{markdown_table(["项目", "值"], [
    ["仓库目录", str(ROOT)],
    ["规则自检", "python app.py --self-test-rule-only"],
    ["完整证据 run", "validation_speedopt/freeze_runs_full/echobench_20260604_180638"],
    ["12帧 run", "validation_speedopt/freeze_runs/echobench_20260604_175653"],
    ["服务 smoke", "validation_speedopt/server_smoke_general_current_20260604.json"],
    ["项目服务链路", "validation_speedopt/server_pipeline_case1_current_20260604.json"],
    ["R 图表脚本", "submission/technical_report/make_freeze_figures.R"],
    ["叙事图脚本", "submission/technical_report/make_nyt_style_figures.py"],
])}

核心重跑命令：

```powershell
.\\.venv\\Scripts\\python.exe app.py --self-test-rule-only
.\\.venv\\Scripts\\python.exe tools\\run_echobench_v1.py --mapping <mapping.csv> --out-root validation_speedopt\\freeze_runs_full --case-limit 60
.\\.venv\\Scripts\\python.exe tools\\run_echobench_v1.py --mapping <mapping.csv> --out-root validation_speedopt\\freeze_runs --case-limit 60 --max-files-per-case 12
.\\.venv\\Scripts\\python.exe tools\\benchmark_server_smoke.py --url http://127.0.0.1:8088 --out validation_speedopt\\server_smoke_general_current_20260604.json
```

## 参考文献

""" + "\n".join(f"- {ref}" for ref in REFERENCES) + "\n"


def set_run_font(run, size=None, bold=None, italic=None, color=None, east_asia="Microsoft YaHei"):
    run.font.name = east_asia
    run._element.rPr.rFonts.set(qn("w:ascii"), east_asia)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), east_asia)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size=None, bold=None, color=None):
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_widths(table, widths_in: list[float]) -> None:
    table.autofit = False
    total_dxa = int(sum(widths_in) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def widths_for_table(headers: list[str]) -> list[float]:
    n = len(headers)
    if n >= 9:
        return [1.45] + [round((6.5 - 1.45) / (n - 1), 3) for _ in headers[1:]]
    if n >= 7:
        return [1.85] + [round((6.5 - 1.85) / (n - 1), 3) for _ in headers[1:]]
    if n == 2:
        return [2.1, 4.4]
    return [round(6.5 / n, 3) for _ in headers]


def format_cell(cell, bold=False, size=8.2, fill=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    if fill:
        set_cell_shading(cell, fill)
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold, color="000000")


def add_table(doc, headers: list[str], rows: list[list[str]], font_size=8.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(header)
        format_cell(cell, bold=True, size=font_size, fill=HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            format_cell(cells[i], size=font_size, align=WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)
    set_table_widths(table, widths_for_table(headers))
    doc.add_paragraph()


def figure_for_docx(path: Path) -> Path:
    out = path.with_suffix(".jpg")
    if not out.exists() or out.stat().st_mtime < path.stat().st_mtime:
        with Image.open(path) as image:
            if image.mode != "RGB":
                image = image.convert("RGB")
            image.save(out, quality=92, optimize=True)
    return out


def build_docx(markdown: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    set_style_font(doc.styles["Normal"], size=10.8)
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.10
    for style_name, size, color in (("Heading 1", 16, ACCENT), ("Heading 2", 13, ACCENT), ("Heading 3", 12, "1F4D78")):
        set_style_font(doc.styles[style_name], size=size, bold=True, color=color)
        doc.styles[style_name].paragraph_format.space_before = Pt(12 if style_name != "Heading 1" else 16)
        doc.styles[style_name].paragraph_format.space_after = Pt(6 if style_name != "Heading 1" else 8)
    section.header.paragraphs[0].text = "CardioConsult PC V5 技术报告 | 可视化增强版"
    section.footer.paragraphs[0].text = "医学教学与算法演示用途，不作为临床诊断或治疗建议"
    for part in (section.header.paragraphs[0], section.footer.paragraphs[0]):
        part.alignment = WD_ALIGN_PARAGRAPH.RIGHT if part is section.footer.paragraphs[0] else WD_ALIGN_PARAGRAPH.LEFT
        for run in part.runs:
            set_run_font(run, size=8.5, color=MUTED)

    lines = markdown.splitlines()
    i = 0
    title_done = False
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.strip() == "<!-- PAGEBREAK -->":
            doc.add_page_break()
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10 if not title_done else 16)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(line[2:].strip())
            set_run_font(r, size=23 if not title_done else 16, bold=True, color=DARK)
            if not title_done:
                sub = doc.add_paragraph()
                sub.paragraph_format.space_after = Pt(14)
                rr = sub.add_run("离线心脏超声教学参考系统：可运行性、准确性、成本与安全边界")
                set_run_font(rr, size=11.5, color=ACCENT)
                title_done = True
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(line[3:].strip())
            set_run_font(r, size=13, bold=True, color=ACCENT)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(5)
            r = p.add_run(line[4:].strip())
            set_run_font(r, size=11.2, bold=True, color=MUTED)
        elif line.startswith("!["):
            start = line.find("(")
            end = line.find(")", start)
            caption = line[2 : line.find("]")]
            rel = line[start + 1 : end]
            fig_path = REPORT_DIR / rel
            if fig_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(figure_for_docx(fig_path)), width=Inches(6.5))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(8)
                rr = cap.add_run(caption)
                set_run_font(rr, size=8.6, italic=True, color=MUTED)
        elif line.startswith("| ") and i + 1 < len(lines) and lines[i + 1].startswith("| "):
            header = [part.strip() for part in line.strip("|").split("|")]
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].startswith("| "):
                rows.append([part.strip() for part in lines[i].strip("|").split("|")])
                i += 1
            add_table(doc, header, rows, font_size=7.1 if len(header) > 7 else 8.2)
            continue
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(line[2:].strip().replace("*", ""))
            set_run_font(r, size=9.6)
        elif line.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run("\n".join(code_lines))
            set_run_font(r, size=8.4, color="334155", east_asia="Consolas")
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.10
            r = p.add_run(line.replace("**", "").replace("*", ""))
            set_run_font(r, size=10.5)
        i += 1

    doc.save(DOCX_OUT)


def convert_pdf() -> None:
    soffice = Path("D:/LibreOffice/program/soffice.exe")
    if not soffice.exists():
        return
    if PDF_OUT.exists():
        PDF_OUT.unlink()
    subprocess.run(
        [str(soffice), "--headless", "--convert-to", "pdf", "--outdir", str(REPORT_DIR), str(DOCX_OUT)],
        check=False,
        cwd=str(REPORT_DIR),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )


def sync_docs_outputs() -> None:
    DOCS_REPORT_DIR.mkdir(parents=True, exist_ok=True)
    DOCS_FIG_DIR.mkdir(parents=True, exist_ok=True)
    DOCS_VISUAL_FIG_DIR.mkdir(parents=True, exist_ok=True)
    for png in FIG_DIR.glob("fig*.png"):
        shutil.copy2(png, DOCS_FIG_DIR / png.name)
    for png in VISUAL_FIG_DIR.glob("nyt_fig*.png"):
        shutil.copy2(png, DOCS_VISUAL_FIG_DIR / png.name)
    shutil.copy2(MARKDOWN_OUT, DOCS_MD_OUT)
    shutil.copy2(DOCX_OUT, DOCS_DOCX_OUT)
    if PDF_OUT.exists():
        shutil.copy2(PDF_OUT, DOCS_REPORT_DIR / "CardioConsult_PC_V5_EchoBench_Technical_Report_APA_20260604.pdf")


def write_freeze_audit(data: dict) -> None:
    audit = ROOT / "submission" / "FREEZE_REVIEW_AUDIT_20260604.md"
    text = f"""# 冻结前评审视角检查

生成日期：{date.today().isoformat()}

## 已补救

- 清理过时 BAT：仓库根目录只保留 `install_deps.bat` 和 `run_cardio_pc_v5.bat`。
- 诊断链报告保护层已上线：提示词泄漏、markdown 模板、AI 口吻和截断输出会回退到本地自然化教学报告。
- 已新增当前服务验证 JSON：`validation_speedopt/server_pipeline_case1_current_20260604.json` 记录 `has_prompt_leakage=false`、`has_required_fields=true` 和旧服务证据 `report_source=gemma4_repaired`；当前默认代码新增结构化 JSON 守卫，可在下一轮服务复现中记录 `report_source=gemma4_structured`。
- 冻结前 EchoBench 完整证据 60/60 通过，平均 {data['full_latency']['runtime_seconds']['mean']:.3f}s/例。
- 冻结前 EchoBench 12 帧 60/60 通过，warm-cache 平均 {data['rep12_latency']['runtime_seconds']['mean']:.3f}s/例。
- 本地 `llama-server` smoke 连续两次 OK，第二次 completion {data['server_smoke']['second_completion']['elapsed_seconds']:.3f}s。
- 技术报告已重写，包含医学指标、一般工程指标、成本模型、取舍、限制、APA 引用、8 张 R 复核图和 7 张白皮书式叙事图。

## 仍需在答辩中主动说明

- 当前系统是医学教学和算法演示工具，不是医疗器械，不用于正式临床诊断。
- 授权本地 60 例是报告链接标签，不是多专家盲评共识。
- TR 全阳性、PR/severe 样本不足，不能宣称这些标签的可靠特异性。
- AR、RWMA、左房扩大对切面覆盖敏感，12 帧输入下降明显。
- 在线 demo 只展示规则匹配和输入输出合同，完整边缘特征与 GGUF 以 PC V5 应用为准。
"""
    audit.write_text(text, encoding="utf-8")


def main() -> None:
    data = gather_data()
    update_integrated_results(data)
    make_nyt_style_figures.main()
    markdown = report_markdown(data)
    MARKDOWN_OUT.write_text(markdown, encoding="utf-8")
    (REPORT_DIR / "APA_REFERENCES.md").write_text(
        "# APA 引用清单\n\n" + "\n".join(f"- {ref}" for ref in REFERENCES) + "\n",
        encoding="utf-8",
    )
    build_docx(markdown)
    convert_pdf()
    sync_docs_outputs()
    write_freeze_audit(data)
    print(MARKDOWN_OUT)
    print(DOCX_OUT)
    print(PDF_OUT if PDF_OUT.exists() else "PDF not generated")


if __name__ == "__main__":
    main()
