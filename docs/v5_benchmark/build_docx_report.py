from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


REPORT_DIR = Path(__file__).resolve().parent
ROOT = REPORT_DIR.parents[1]
FIG_DIR = REPORT_DIR / "figures"
FIG_DOCX_DIR = REPORT_DIR / "figures_docx"
OUT = REPORT_DIR / "CardioConsult_PC_V5_EchoBench_Technical_Report_APA_20260604.docx"


ACCENT = "2E74B5"
DARK = "1F4D78"
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "CBD5E1"


def set_run_font(run, size=None, bold=None, italic=None, color=None, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size=None, bold=None, color=None, east_asia="Microsoft YaHei"):
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def format_cell(cell, bold=False, size=8.5, fill=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    if fill:
        set_cell_shading(cell, fill)
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold, color="000000")


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    set_table_borders(table, "D8DEE9")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(text)
    set_run_font(r2, size=10)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None, font_size=8.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    if widths is None:
        widths = [6.3 / len(headers)] * len(headers)

    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_row_cant_split(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.width = Inches(widths[idx])
        cell.text = str(text)
        format_cell(cell, bold=True, size=font_size, fill=HEADER_FILL)

    for row in rows:
        new_row = table.add_row()
        set_row_cant_split(new_row)
        cells = new_row.cells
        for idx, text in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            cells[idx].text = str(text)
            align = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            format_cell(cells[idx], size=font_size, align=align)
    doc.add_paragraph()
    return table


def add_figure(doc, filename, caption):
    path = figure_for_docx(filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.2))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    set_run_font(r, size=9.5, italic=True, color="555555")


def figure_for_docx(filename):
    src = FIG_DIR / filename
    FIG_DOCX_DIR.mkdir(parents=True, exist_ok=True)
    out = FIG_DOCX_DIR / (Path(filename).stem + ".jpg")
    if not out.exists() or out.stat().st_mtime < src.stat().st_mtime:
        with Image.open(src) as im:
            if im.mode != "RGB":
                im = im.convert("RGB")
            im.save(out, quality=94, optimize=True)
    return out


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], size=11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        style = styles[name]
        set_style_font(style, size=size, bold=True, color=color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "CardioConsult PC V5 EchoBench 技术报告"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color="666666")

    footer = section.footer.paragraphs[0]
    footer.text = "医学教学与算法演示用途，不作为临床诊断或治疗建议"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=8.5, color="666666")


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("CardioConsult PC V5 EchoBench 技术报告")
    set_run_font(r, size=22, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("APA 引用版 · 带图带表 · 生成日期：2026-06-04")
    set_run_font(r, size=11, color="555555")

    add_callout(
        doc,
        "技术摘要",
        "CardioConsult PC V5 面向心脏超声教学参考和基层辅助初筛场景，目标是在普通 PC 上离线读取 PNG、DICOM/DCOM 与超声动图文件，输出从大方向到最小病症的结构化疑似诊断文字。V5 在原有 B-mode、Color Doppler、层级标签规则和 Gemma4 4B GGUF 本地报告生成基础上，加入 EchoNet-Dynamic 动态 B-mode 校准层，用于增强 EF 和左室收缩功能减低识别。"
    )

    doc.add_heading("关键结论", level=1)
    bullets = [
        "完整证据场景 60/60 例成功，平均 3.76 秒/例；12 帧代表抽样场景 60/60 例成功，平均 2.56 秒/例。",
        "完整证据下 MR F1 为 0.964，AR F1 为 0.700，低 EF F1 为 0.857；MR 与低 EF 已达到较稳定的教学演示水平。",
        "12 帧场景下 MR F1 仍为 0.936，但 AR F1 降至 0.326，说明主动脉瓣反流强依赖切面覆盖和彩色多普勒代表帧。",
        "GGUF 本地生成建议使用常驻 llama-server：冷启动首次 completion 约 8.78 秒，热启动 completion 约 0.49 秒。"
    ]
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(item)
        set_run_font(r)

    doc.add_heading("完整证据场景：核心指标", level=1)
    add_para(doc, "完整证据场景使用每例所有可用文件，平均 17.3 个文件/例，最多 33 个文件/例。该场景代表资料相对充分时的系统能力上限。TR 在本批数据中没有阴性样本，因此特异性不能作为有效结论。")
    full_rows = [
        ["MR（二尖瓣）", "60", "55", "0.933", "0.964", "0.600", "0.964"],
        ["TR（三尖瓣）", "60", "60", "1.000", "1.000", "0.000*", "1.000"],
        ["AR（主动脉瓣）", "60", "29", "0.700", "0.724", "0.677", "0.700"],
        ["低 EF", "60", "6", "0.967", "1.000", "0.963", "0.857"],
        ["RWMA", "60", "3", "0.967", "0.333", "1.000", "0.500"],
        ["左房扩大", "60", "8", "0.883", "1.000", "0.865", "0.696"],
    ]
    add_table(
        doc,
        ["标签", "n", "阳性金标准", "准确率", "敏感性", "特异性", "F1"],
        full_rows,
        widths=[1.9, 0.45, 0.85, 0.75, 0.75, 0.75, 0.65],
    )
    add_figure(doc, "fig1_f1_full_vs_12frame.png", "图 1. EchoBench v1 F1：完整证据与 12 帧代表抽样对比。")

    doc.add_heading("12 帧代表抽样：输入受限场景", level=1)
    add_para(doc, "12 帧场景将每例文件做均匀代表抽样，而不是只取前 12 个文件。该场景更接近现场演示和标准输入上限，但会牺牲部分细粒度病症定位能力。")
    rep_rows = [
        ["MR（二尖瓣）", "60", "55", "0.883", "0.927", "0.400", "0.936"],
        ["TR（三尖瓣）", "60", "60", "1.000", "1.000", "0.000*", "1.000"],
        ["AR（主动脉瓣）", "60", "29", "0.517", "0.241", "0.774", "0.326"],
        ["低 EF", "60", "6", "0.917", "0.667", "0.944", "0.615"],
        ["RWMA", "60", "3", "0.933", "0.333", "0.965", "0.333"],
        ["左房扩大", "60", "8", "0.750", "0.375", "0.808", "0.286"],
    ]
    add_table(
        doc,
        ["标签", "n", "阳性金标准", "准确率", "敏感性", "特异性", "F1"],
        rep_rows,
        widths=[1.9, 0.45, 0.85, 0.75, 0.75, 0.75, 0.65],
    )
    add_figure(doc, "fig2_accuracy_full_vs_12frame.png", "图 2. EchoBench v1 准确率：完整证据与 12 帧代表抽样对比。")

    doc.add_heading("延迟与本地 GGUF 性能", level=1)
    add_para(doc, "完整证据场景平均 3.76 秒/例，P95 为 5.51 秒；12 帧场景平均 2.56 秒/例，P95 为 3.20 秒。规则与小模型校准流水线不包含每例完整 GGUF 文本生成。")
    add_table(
        doc,
        ["场景", "成功例数", "平均秒/例", "P50", "P90", "P95", "P99", "最大值"],
        [
            ["完整证据", "60/60", "3.761", "3.311", "5.012", "5.513", "6.704", "7.282"],
            ["12 帧代表抽样", "60/60", "2.562", "2.471", "2.796", "3.201", "3.624", "3.680"],
        ],
        widths=[1.25, 0.75, 0.85, 0.65, 0.65, 0.65, 0.65, 0.65],
    )
    add_figure(doc, "fig3_latency_full_vs_12frame.png", "图 3. EchoBench v1 每例延迟百分位。")
    add_table(
        doc,
        ["GGUF/llama.cpp 指标", "结果"],
        [
            ["GGUF 文件", "gemma-4-4b-it-Q4_K_M.gguf"],
            ["SHA256", "519b9793ed6ce0ff530f1b7c96e848e08e49e7af4d57bb97f76215963a54146d"],
            ["llama.cpp build", "b9469"],
            ["CPU threads", "14"],
            ["prompt processing", "37.76 tokens/s"],
            ["generation", "6.19 tokens/s"],
            ["server 首次 completion", "8.775 s"],
            ["server 热启动 completion", "0.492 s"],
        ],
        widths=[2.2, 4.1],
        font_size=8.1,
    )

    doc.add_heading("数据、模型与方法", level=1)
    add_para(doc, "本次主测试集来自授权本地 DICOM/报告时间映射，共 60 个病例。该数据只用于本地授权教育验证，不随代码发布，不作为公开数据集再分发。报告链接标签被用作当前 benchmark 的报告链接金标准，但不是独立多专家盲法复核金标准。")
    add_para(doc, "V5 使用 EchoNet-Dynamic 公开数据进行动态 B-mode EF 校准。实际读取 10,030 个心尖四腔 .avi 视频、FileList.csv 中的 EF/ESV/EDV/FPS/帧数/官方 split，以及 VolumeTracings.csv 中的专家左室追踪帧。EchoNet-Dynamic 原论文证明心超视频深度学习可用于逐搏心功能评估（Ouyang et al., 2020）。")
    add_table(
        doc,
        ["子任务", "选中模型", "选择依据"],
        [
            ["EF 回归", "HistGradientBoostingRegressor", "验证集 MAE/RMSE 优于 Ridge 与 MLP"],
            ["低 EF 分类", "LogisticRegression", "验证集 F1 最优，AUC 接近树模型，推理更轻"],
        ],
        widths=[1.35, 2.25, 2.7],
    )
    add_table(
        doc,
        ["EchoNet-Dynamic held-out 指标", "数值"],
        [
            ["EF MAE", "7.271"],
            ["EF RMSE", "9.603"],
            ["EF 相关系数", "0.647"],
            ["低 EF 准确率", "0.770"],
            ["低 EF 精确率", "0.479"],
            ["低 EF 召回率", "0.515"],
            ["低 EF F1", "0.496"],
            ["低 EF AUC", "0.764"],
        ],
        widths=[3.15, 1.3],
    )
    add_figure(doc, "fig4_echonet_training_metrics.png", "图 4. V5 EchoNet-Dynamic 校准层 held-out 指标。")

    doc.add_heading("Benchmark 设计与现有方案比较", level=1)
    add_para(doc, "EchoBench v1 是项目级 benchmark，而不是单脚本 smoke test。它包含单例交互、离线批量和常驻 server 三类场景。本次主报告使用离线批量完整证据、离线批量 12 帧代表抽样和 GGUF server smoke。性能测量思路参考 MLPerf 对离线、客户端和服务端场景的拆分，但本报告不是 MLPerf 官方提交结果（MLCommons, n.d.-a, n.d.-b）。")
    add_para(doc, "现有公开心超 AI 方案中，EchoNet-Dynamic 聚焦 EF 与左室功能，CAMUS 聚焦 2D 心超分割（Leclerc et al., 2019; Ouyang et al., 2020）。CardioConsult V5 的差异化在于多格式输入、B-mode 与 Color Doppler 代理特征、动图兼容、层级病症输出和本地离线 Gemma4 4B GGUF 中文报告生成。")

    doc.add_heading("成本模型与工程取舍", level=1)
    add_para(doc, "V5 的成本结构主要是一次性硬件和模型文件成本，运行时不需要按病例调用云 API。完整证据场景按平均 3.76 秒/例估算，规则流水线理论吞吐约 957 例/小时；12 帧场景按平均 2.56 秒/例估算，理论吞吐约 1,405 例/小时。真实 UI 使用会受人工选文件、磁盘读取、DICOM 解码和 GGUF 文本生成影响，因此演示吞吐应按更保守数值估计。")
    add_para(doc, "端到端大型视频模型可能在 EF 或分割任务上更强，但会带来 GPU/NPU 依赖、模型体积、训练时间和移动端迁移成本。V5 选择小模型校准和规则融合，是为了确保普通 PC 可跑、移动端后续可迁移，并且每条诊断都有特征证据和规则路径可追踪。")

    doc.add_heading("限制与下一步", level=1)
    limits = [
        "本地 60 例测试集规模有限，且来自授权本地数据，不是公开多中心盲法验证集。",
        "报告链接标签是当前 benchmark 的金标准，但不是 2-3 位超声专家独立盲评后的共识标签。",
        "TR 在本批数据中全为阳性，PR 和 severe 标签样本不足，不能做有效结论。",
        "EchoNet-Dynamic 只增强 A4C B-mode 心功能相关任务，不提供瓣膜反流分级监督。",
        "AR、RWMA 和左房扩大在 12 帧抽样下降明显，未来必须增强切面识别、动图分割和标签平衡。",
    ]
    for item in limits:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(item)
        set_run_font(r)
    add_callout(
        doc,
        "安全边界",
        "本系统仅用于医学教学、算法演示和基层参考，不是医疗器械输出，不作为正式临床诊断、治疗建议或医嘱。若患者存在胸痛、晕厥、明显呼吸困难、低血压、发绀或急性心衰表现，应直接进入正式医疗流程并由有资质医师复核。"
    )

    doc.add_heading("可复现信息", level=1)
    add_table(
        doc,
        ["项目", "值"],
        [
            ["V5 根目录", "<V5_ROOT>"],
            ["PC 应用目录", "<V5_ROOT>\\05_pc_v5"],
            ["完整证据 run", "runs\\echobench_20260604_114319"],
            ["12 帧 run", "runs\\echobench_20260604_114016"],
            ["EchoNet 训练报告", "training\\echonet_v5\\training_report.json"],
            ["R 图表脚本", "reports\\make_benchmark_figures.R"],
        ],
        widths=[2.0, 4.3],
        font_size=8.2,
    )

    doc.add_heading("参考文献", level=1)
    refs = [
        "FDA, Health Canada, & Medicines and Healthcare products Regulatory Agency. (2021). Good machine learning practice for medical device development: Guiding principles. https://www.gov.uk/government/publications/good-machine-learning-practice-for-medical-device-development-guiding-principles",
        "ggml-org. (n.d.). GGUF file format - llama.cpp. Retrieved June 4, 2026, from https://www.mintlify.com/ggml-org/llama.cpp/concepts/gguf-format",
        "Google AI for Developers. (n.d.). Get started with Gemma models. Retrieved June 4, 2026, from https://ai.google.dev/gemma/docs/get_started",
        "Lang, R. M., Badano, L. P., Mor-Avi, V., Afilalo, J., Armstrong, A., Ernande, L., Flachskampf, F. A., Foster, E., Goldstein, S. A., Kuznetsova, T., Lancellotti, P., Muraru, D., Picard, M. H., Rietzschel, E. R., Rudski, L., Spencer, K. T., Tsang, W., & Voigt, J.-U. (2015). Recommendations for cardiac chamber quantification by echocardiography in adults: An update from the American Society of Echocardiography and the European Association of Cardiovascular Imaging. Journal of the American Society of Echocardiography, 28(1), 1-39.e14. https://doi.org/10.1016/j.echo.2014.10.003",
        "Leclerc, S., Smistad, E., Pedrosa, J., Ostvik, A., Cervenansky, F., Espinosa, F., Espeland, T., Berg, E. A. R., Jodoin, P.-M., Grenier, T., Lartizien, C., Dhooge, J., Lovstakken, L., Bernard, O., & Grenier, T. (2019). Deep learning for segmentation using an open large-scale dataset in 2D echocardiography. IEEE Transactions on Medical Imaging, 38(9), 2198-2210. https://doi.org/10.1109/TMI.2019.2900516",
        "MLCommons. (n.d.-a). MLPerf Client. Retrieved June 4, 2026, from https://mlcommons.org/benchmarks/client/",
        "MLCommons. (n.d.-b). MLPerf Inference benchmarks. Retrieved June 4, 2026, from https://docs.mlcommons.org/inference/",
        "MLCommons. (n.d.-c). MedPerf: An open benchmarking platform for medical artificial intelligence using federated evaluation. Retrieved June 4, 2026, from https://github.com/mlcommons/medperf",
        "OpenAI. (2025). Introducing HealthBench. https://openai.com/index/healthbench/",
        "Ouyang, D., He, B., Ghorbani, A., Yuan, N., Ebinger, J., Langlotz, C. P., Heidenreich, P. A., Harrington, R. A., Liang, D. H., Ashley, E. A., & Zou, J. Y. (2020). Video-based AI for beat-to-beat assessment of cardiac function. Nature, 580(7802), 252-256. https://doi.org/10.1038/s41586-020-2145-8",
        "Zoghbi, W. A., Adams, D., Bonow, R. O., Enriquez-Sarano, M., Foster, E., Grayburn, P. A., Hahn, R. T., Han, Y., Hung, J., Lang, R. M., Little, S. H., Shah, D. J., Shernan, S., Thavendiranathan, P., Thomas, J. D., & Weissman, N. J. (2017). Recommendations for noninvasive evaluation of native valvular regurgitation: A report from the American Society of Echocardiography developed in collaboration with the Society for Cardiovascular Magnetic Resonance. Journal of the American Society of Echocardiography, 30(4), 303-371. https://doi.org/10.1016/j.echo.2017.01.007",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(ref)
        set_run_font(r, size=9.4)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
